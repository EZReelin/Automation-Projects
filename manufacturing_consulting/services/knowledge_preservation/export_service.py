"""
Export service for SOPs and knowledge documents.

Supports export to multiple formats: PDF, Word, Markdown, HTML.
"""

import io
from datetime import datetime
from typing import BinaryIO

from jinja2 import Environment, BaseLoader
from sqlalchemy.ext.asyncio import AsyncSession

from models.knowledge_preservation import SOP, KnowledgeDomain
from utils.logging import ServiceLogger


class ExportService:
    """
    Service for exporting SOPs to various formats.
    
    Supports:
    - PDF export
    - Word (DOCX) export
    - Markdown export
    - HTML export
    """
    
    def __init__(self, session: AsyncSession, tenant_id: str):
        self.session = session
        self.tenant_id = tenant_id
        self.logger = ServiceLogger("export")
        self.jinja_env = Environment(loader=BaseLoader())
    
    async def export_sop_to_markdown(self, sop: SOP) -> str:
        """
        Export SOP to Markdown format.
        
        Args:
            sop: SOP to export
            
        Returns:
            Markdown formatted string
        """
        content = sop.content
        
        md_parts = [
            f"# {sop.title}",
            "",
            f"**SOP Number:** {sop.sop_number}",
            f"**Version:** {sop.version}",
            f"**Effective Date:** {sop.effective_date.strftime('%Y-%m-%d') if sop.effective_date else 'TBD'}",
            f"**Status:** {sop.status.value.replace('_', ' ').title()}",
            "",
            "---",
            "",
        ]
        
        # Purpose
        if content.get("purpose"):
            md_parts.extend([
                "## Purpose",
                "",
                content["purpose"],
                "",
            ])
        
        # Scope
        if content.get("scope"):
            md_parts.extend([
                "## Scope",
                "",
                content["scope"],
                "",
            ])
        
        # Responsibilities
        if content.get("responsibilities"):
            md_parts.extend([
                "## Responsibilities",
                "",
            ])
            for resp in content["responsibilities"]:
                md_parts.append(f"- {resp}")
            md_parts.append("")
        
        # Definitions
        if content.get("definitions"):
            md_parts.extend([
                "## Definitions",
                "",
            ])
            for term, definition in content["definitions"].items():
                md_parts.append(f"- **{term}:** {definition}")
            md_parts.append("")
        
        # Safety Requirements
        if content.get("safety_requirements"):
            md_parts.extend([
                "## Safety Requirements",
                "",
                "⚠️ **Important Safety Information**",
                "",
            ])
            for req in content["safety_requirements"]:
                md_parts.append(f"- {req}")
            md_parts.append("")
        
        # Equipment Needed
        if content.get("equipment_needed"):
            md_parts.extend([
                "## Equipment Needed",
                "",
            ])
            for equip in content["equipment_needed"]:
                md_parts.append(f"- {equip}")
            md_parts.append("")
        
        # Materials Needed
        if content.get("materials_needed"):
            md_parts.extend([
                "## Materials Needed",
                "",
            ])
            for mat in content["materials_needed"]:
                md_parts.append(f"- {mat}")
            md_parts.append("")
        
        # Procedure Steps
        if content.get("procedure_steps"):
            md_parts.extend([
                "## Procedure",
                "",
            ])
            for step in content["procedure_steps"]:
                step_num = step.get("step_number", "")
                title = step.get("title", "")
                desc = step.get("description", "")
                
                md_parts.append(f"### Step {step_num}: {title}")
                md_parts.append("")
                md_parts.append(desc)
                md_parts.append("")
                
                # Cautions
                if step.get("cautions"):
                    md_parts.append("**⚠️ Cautions:**")
                    for caution in step["cautions"]:
                        md_parts.append(f"- {caution}")
                    md_parts.append("")
                
                # Notes
                if step.get("notes"):
                    md_parts.append("**📝 Notes:**")
                    for note in step["notes"]:
                        md_parts.append(f"- {note}")
                    md_parts.append("")
        
        # Quality Checkpoints
        if content.get("quality_checkpoints"):
            md_parts.extend([
                "## Quality Checkpoints",
                "",
            ])
            for checkpoint in content["quality_checkpoints"]:
                md_parts.append(f"- [ ] {checkpoint}")
            md_parts.append("")
        
        # Troubleshooting
        if content.get("troubleshooting"):
            md_parts.extend([
                "## Troubleshooting",
                "",
                "| Problem | Cause | Solution |",
                "|---------|-------|----------|",
            ])
            for item in content["troubleshooting"]:
                problem = item.get("problem", "")
                cause = item.get("cause", "")
                solution = item.get("solution", "")
                md_parts.append(f"| {problem} | {cause} | {solution} |")
            md_parts.append("")
        
        # References
        if content.get("references"):
            md_parts.extend([
                "## References",
                "",
            ])
            for ref in content["references"]:
                md_parts.append(f"- {ref}")
            md_parts.append("")
        
        # Footer
        md_parts.extend([
            "---",
            "",
            f"*Generated on {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}*",
        ])
        
        return "\n".join(md_parts)
    
    async def export_sop_to_html(self, sop: SOP) -> str:
        """
        Export SOP to HTML format.
        
        Args:
            sop: SOP to export
            
        Returns:
            HTML formatted string
        """
        markdown_content = await self.export_sop_to_markdown(sop)
        
        # Convert markdown to HTML using simple conversion
        # In production, use a proper markdown library
        html = self._markdown_to_html(markdown_content)
        
        # Wrap in HTML document
        html_doc = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{sop.title} - {sop.sop_number}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 2rem;
            color: #333;
        }}
        h1 {{ color: #1a73e8; border-bottom: 2px solid #1a73e8; padding-bottom: 0.5rem; }}
        h2 {{ color: #2c3e50; margin-top: 2rem; }}
        h3 {{ color: #34495e; }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1rem 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 0.75rem;
            text-align: left;
        }}
        th {{ background-color: #f5f5f5; }}
        .caution {{
            background-color: #fff3cd;
            border-left: 4px solid #ffc107;
            padding: 1rem;
            margin: 1rem 0;
        }}
        .note {{
            background-color: #e7f3ff;
            border-left: 4px solid #1a73e8;
            padding: 1rem;
            margin: 1rem 0;
        }}
        code {{ background-color: #f5f5f5; padding: 0.2rem 0.4rem; border-radius: 3px; }}
        hr {{ border: none; border-top: 1px solid #ddd; margin: 2rem 0; }}
        ul {{ list-style-type: disc; padding-left: 1.5rem; }}
        .checkbox {{ list-style-type: none; }}
        .checkbox li::before {{ content: "☐ "; }}
        @media print {{
            body {{ max-width: none; }}
            h1, h2 {{ page-break-after: avoid; }}
        }}
    </style>
</head>
<body>
{html}
</body>
</html>"""
        
        return html_doc
    
    async def export_sop_to_docx(self, sop: SOP) -> BinaryIO:
        """
        Export SOP to Word (DOCX) format.
        
        Args:
            sop: SOP to export
            
        Returns:
            BytesIO object containing DOCX file
        """
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        content = sop.content
        
        # Title
        title = doc.add_heading(sop.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata table
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        metadata = [
            ("SOP Number", sop.sop_number),
            ("Version", sop.version),
            ("Effective Date", sop.effective_date.strftime('%Y-%m-%d') if sop.effective_date else 'TBD'),
            ("Status", sop.status.value.replace('_', ' ').title()),
        ]
        
        for i, (label, value) in enumerate(metadata):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        doc.add_paragraph()
        
        # Purpose
        if content.get("purpose"):
            doc.add_heading("Purpose", level=1)
            doc.add_paragraph(content["purpose"])
        
        # Scope
        if content.get("scope"):
            doc.add_heading("Scope", level=1)
            doc.add_paragraph(content["scope"])
        
        # Responsibilities
        if content.get("responsibilities"):
            doc.add_heading("Responsibilities", level=1)
            for resp in content["responsibilities"]:
                doc.add_paragraph(resp, style='List Bullet')
        
        # Definitions
        if content.get("definitions"):
            doc.add_heading("Definitions", level=1)
            for term, definition in content["definitions"].items():
                p = doc.add_paragraph()
                p.add_run(f"{term}: ").bold = True
                p.add_run(definition)
        
        # Safety Requirements
        if content.get("safety_requirements"):
            doc.add_heading("Safety Requirements", level=1)
            for req in content["safety_requirements"]:
                doc.add_paragraph(req, style='List Bullet')
        
        # Equipment Needed
        if content.get("equipment_needed"):
            doc.add_heading("Equipment Needed", level=1)
            for equip in content["equipment_needed"]:
                doc.add_paragraph(equip, style='List Bullet')
        
        # Materials Needed
        if content.get("materials_needed"):
            doc.add_heading("Materials Needed", level=1)
            for mat in content["materials_needed"]:
                doc.add_paragraph(mat, style='List Bullet')
        
        # Procedure Steps
        if content.get("procedure_steps"):
            doc.add_heading("Procedure", level=1)
            for step in content["procedure_steps"]:
                step_num = step.get("step_number", "")
                title = step.get("title", "")
                desc = step.get("description", "")
                
                doc.add_heading(f"Step {step_num}: {title}", level=2)
                doc.add_paragraph(desc)
                
                if step.get("cautions"):
                    p = doc.add_paragraph()
                    p.add_run("Cautions:").bold = True
                    for caution in step["cautions"]:
                        doc.add_paragraph(f"⚠️ {caution}", style='List Bullet')
                
                if step.get("notes"):
                    p = doc.add_paragraph()
                    p.add_run("Notes:").bold = True
                    for note in step["notes"]:
                        doc.add_paragraph(f"📝 {note}", style='List Bullet')
        
        # Quality Checkpoints
        if content.get("quality_checkpoints"):
            doc.add_heading("Quality Checkpoints", level=1)
            for checkpoint in content["quality_checkpoints"]:
                doc.add_paragraph(f"☐ {checkpoint}")
        
        # Troubleshooting
        if content.get("troubleshooting"):
            doc.add_heading("Troubleshooting", level=1)
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            headers = table.rows[0].cells
            headers[0].text = "Problem"
            headers[1].text = "Cause"
            headers[2].text = "Solution"
            
            for item in content["troubleshooting"]:
                row = table.add_row().cells
                row[0].text = item.get("problem", "")
                row[1].text = item.get("cause", "")
                row[2].text = item.get("solution", "")
        
        # References
        if content.get("references"):
            doc.add_heading("References", level=1)
            for ref in content["references"]:
                doc.add_paragraph(ref, style='List Bullet')
        
        # Footer
        doc.add_paragraph()
        doc.add_paragraph(f"Generated on {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
        
        # Save to BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    async def export_sop_to_pdf(self, sop: SOP) -> BinaryIO:
        """
        Export SOP to PDF format.
        
        Args:
            sop: SOP to export
            
        Returns:
            BytesIO object containing PDF file
        """
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            PageBreak, ListFlowable, ListItem
        )
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
        )
        
        styles = getSampleStyleSheet()
        
        # Custom styles
        styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            textColor=colors.HexColor('#1a73e8'),
        ))
        
        styles.add(ParagraphStyle(
            name='SectionHeader',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=20,
            spaceAfter=10,
            textColor=colors.HexColor('#2c3e50'),
        ))
        
        story = []
        content = sop.content
        
        # Title
        story.append(Paragraph(sop.title, styles['CustomTitle']))
        
        # Metadata table
        meta_data = [
            ['SOP Number:', sop.sop_number],
            ['Version:', sop.version],
            ['Effective Date:', sop.effective_date.strftime('%Y-%m-%d') if sop.effective_date else 'TBD'],
            ['Status:', sop.status.value.replace('_', ' ').title()],
        ]
        meta_table = Table(meta_data, colWidths=[1.5*inch, 3*inch])
        meta_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(meta_table)
        story.append(Spacer(1, 20))
        
        # Purpose
        if content.get("purpose"):
            story.append(Paragraph("Purpose", styles['SectionHeader']))
            story.append(Paragraph(content["purpose"], styles['Normal']))
        
        # Scope
        if content.get("scope"):
            story.append(Paragraph("Scope", styles['SectionHeader']))
            story.append(Paragraph(content["scope"], styles['Normal']))
        
        # Responsibilities
        if content.get("responsibilities"):
            story.append(Paragraph("Responsibilities", styles['SectionHeader']))
            items = [ListItem(Paragraph(r, styles['Normal'])) for r in content["responsibilities"]]
            story.append(ListFlowable(items, bulletType='bullet'))
        
        # Safety Requirements
        if content.get("safety_requirements"):
            story.append(Paragraph("Safety Requirements", styles['SectionHeader']))
            items = [ListItem(Paragraph(f"⚠️ {r}", styles['Normal'])) for r in content["safety_requirements"]]
            story.append(ListFlowable(items, bulletType='bullet'))
        
        # Equipment Needed
        if content.get("equipment_needed"):
            story.append(Paragraph("Equipment Needed", styles['SectionHeader']))
            items = [ListItem(Paragraph(e, styles['Normal'])) for e in content["equipment_needed"]]
            story.append(ListFlowable(items, bulletType='bullet'))
        
        # Procedure Steps
        if content.get("procedure_steps"):
            story.append(Paragraph("Procedure", styles['SectionHeader']))
            for step in content["procedure_steps"]:
                step_num = step.get("step_number", "")
                title = step.get("title", "")
                desc = step.get("description", "")
                
                story.append(Paragraph(f"Step {step_num}: {title}", styles['Heading3']))
                story.append(Paragraph(desc, styles['Normal']))
                
                if step.get("cautions"):
                    story.append(Paragraph("<b>Cautions:</b>", styles['Normal']))
                    items = [ListItem(Paragraph(c, styles['Normal'])) for c in step["cautions"]]
                    story.append(ListFlowable(items, bulletType='bullet'))
                
                story.append(Spacer(1, 10))
        
        # Quality Checkpoints
        if content.get("quality_checkpoints"):
            story.append(Paragraph("Quality Checkpoints", styles['SectionHeader']))
            items = [ListItem(Paragraph(f"☐ {c}", styles['Normal'])) for c in content["quality_checkpoints"]]
            story.append(ListFlowable(items, bulletType='bullet'))
        
        # Troubleshooting
        if content.get("troubleshooting"):
            story.append(Paragraph("Troubleshooting", styles['SectionHeader']))
            
            ts_data = [['Problem', 'Cause', 'Solution']]
            for item in content["troubleshooting"]:
                ts_data.append([
                    item.get("problem", ""),
                    item.get("cause", ""),
                    item.get("solution", ""),
                ])
            
            ts_table = Table(ts_data, colWidths=[2*inch, 2*inch, 2*inch])
            ts_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f5f5f5')),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#dddddd')),
                ('PADDING', (0, 0), (-1, -1), 8),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            story.append(ts_table)
        
        # Footer
        story.append(Spacer(1, 30))
        story.append(Paragraph(
            f"<i>Generated on {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}</i>",
            styles['Normal']
        ))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        return buffer
    
    def _markdown_to_html(self, markdown: str) -> str:
        """Simple markdown to HTML conversion."""
        import re
        
        html = markdown
        
        # Headers
        html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
        html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
        html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
        
        # Bold
        html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
        
        # Italic
        html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)
        
        # Lists
        html = re.sub(r'^- \[ \] (.+)$', r'<li class="checkbox">\1</li>', html, flags=re.MULTILINE)
        html = re.sub(r'^- (.+)$', r'<li>\1</li>', html, flags=re.MULTILINE)
        
        # Wrap consecutive li tags in ul
        html = re.sub(r'(<li[^>]*>.*?</li>\n?)+', r'<ul>\g<0></ul>', html)
        
        # Tables (basic)
        lines = html.split('\n')
        in_table = False
        new_lines = []
        
        for line in lines:
            if '|' in line and not in_table:
                in_table = True
                new_lines.append('<table>')
            
            if in_table:
                if '|' not in line:
                    in_table = False
                    new_lines.append('</table>')
                    new_lines.append(line)
                elif '---' not in line:
                    cells = [c.strip() for c in line.split('|')[1:-1]]
                    row = '<tr>' + ''.join(f'<td>{c}</td>' for c in cells) + '</tr>'
                    new_lines.append(row)
            else:
                new_lines.append(line)
        
        if in_table:
            new_lines.append('</table>')
        
        html = '\n'.join(new_lines)
        
        # Horizontal rules
        html = re.sub(r'^---$', r'<hr>', html, flags=re.MULTILINE)
        
        # Paragraphs
        html = re.sub(r'\n\n', r'</p><p>', html)
        html = f'<p>{html}</p>'
        
        # Clean up empty paragraphs
        html = re.sub(r'<p>\s*</p>', '', html)
        html = re.sub(r'<p>(<h[123]>)', r'\1', html)
        html = re.sub(r'(</h[123]>)</p>', r'\1', html)
        
        return html
