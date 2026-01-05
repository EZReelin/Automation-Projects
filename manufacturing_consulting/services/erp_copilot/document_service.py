"""
ERP Document management service.

Handles document ingestion, processing, and indexing for RAG retrieval.
"""

from datetime import datetime
from typing import Any, BinaryIO
from uuid import uuid4

from sqlalchemy import and_, func, select
from sqlalchemy.ext.asyncio import AsyncSession

from models.erp_copilot import (
    ERPConfiguration, ERPDocument, ERPDocumentChunk,
    ERPSystem, DocumentType
)
from utils.logging import ServiceLogger
from utils.ai_client import ai_client
from utils.vector_search import tenant_vector_store, text_chunker
from config.settings import settings


class ERPDocumentService:
    """
    Service for managing ERP documentation.
    
    Provides:
    - Document ingestion from files and URLs
    - Content extraction and cleaning
    - Chunking and embedding generation
    - Semantic search over documents
    """
    
    def __init__(self, session: AsyncSession, tenant_id: str):
        self.session = session
        self.tenant_id = tenant_id
        self.logger = ServiceLogger("erp_document")
    
    async def configure_erp(
        self,
        erp_system: ERPSystem,
        system_version: str | None = None,
        instance_name: str | None = None,
        enabled_modules: list[str] | None = None,
        custom_terminology: dict | None = None,
    ) -> ERPConfiguration:
        """
        Create or update ERP configuration for tenant.
        
        Args:
            erp_system: ERP system type
            system_version: Version string
            instance_name: Instance identifier
            enabled_modules: List of enabled modules
            custom_terminology: Custom term mappings
            
        Returns:
            ERPConfiguration instance
        """
        # Check for existing config
        existing = await self._get_erp_config()
        
        if existing:
            existing.erp_system = erp_system
            existing.system_version = system_version
            existing.instance_name = instance_name
            if enabled_modules:
                existing.enabled_modules = enabled_modules
            if custom_terminology:
                existing.custom_terminology = custom_terminology
            existing.updated_at = datetime.utcnow()
            await self.session.flush()
            return existing
        
        config = ERPConfiguration(
            id=str(uuid4()),
            tenant_id=self.tenant_id,
            erp_system=erp_system,
            system_version=system_version,
            instance_name=instance_name,
            enabled_modules=enabled_modules or [],
            custom_terminology=custom_terminology or {},
        )
        
        self.session.add(config)
        await self.session.flush()
        
        return config
    
    async def ingest_document(
        self,
        title: str,
        content: str,
        document_type: DocumentType = DocumentType.OTHER,
        module: str | None = None,
        source_url: str | None = None,
        erp_menu_path: str | None = None,
        tags: list[str] | None = None,
        **kwargs: Any,
    ) -> ERPDocument:
        """
        Ingest a document into the system.
        
        Args:
            title: Document title
            content: Document content
            document_type: Type of documentation
            module: ERP module this relates to
            source_url: Original source URL
            erp_menu_path: Menu path in ERP
            tags: Document tags
            **kwargs: Additional attributes
            
        Returns:
            Created ERPDocument instance
        """
        self.logger.log_operation_start(
            "ingest_document",
            tenant_id=self.tenant_id,
            title=title[:50],
        )
        
        # Get ERP config
        config = await self._get_erp_config()
        if not config:
            raise ValueError("ERP configuration not found. Configure ERP first.")
        
        # Clean content
        cleaned_content = await self._clean_content(content)
        
        # Generate summary
        summary = await self._generate_summary(cleaned_content)
        
        # Extract keywords
        keywords = await self._extract_keywords(cleaned_content)
        
        document = ERPDocument(
            id=str(uuid4()),
            tenant_id=self.tenant_id,
            erp_config_id=config.id,
            title=title,
            document_type=document_type,
            content_raw=content,
            content_cleaned=cleaned_content,
            content_summary=summary,
            module=module,
            source_url=source_url,
            erp_menu_path=erp_menu_path,
            tags=tags or [],
            keywords=keywords,
            is_processed=False,
            **kwargs,
        )
        
        self.session.add(document)
        await self.session.flush()
        
        # Process and index document
        await self._process_document(document)
        
        self.logger.log_operation_complete(
            "ingest_document",
            tenant_id=self.tenant_id,
            document_id=document.id,
        )
        
        return document
    
    async def ingest_document_file(
        self,
        file: BinaryIO,
        filename: str,
        document_type: DocumentType = DocumentType.OTHER,
        **kwargs: Any,
    ) -> ERPDocument:
        """
        Ingest a document from a file.
        
        Supports PDF, DOCX, TXT, MD files.
        
        Args:
            file: File object
            filename: Original filename
            document_type: Type of documentation
            **kwargs: Additional attributes
            
        Returns:
            Created ERPDocument instance
        """
        # Extract content based on file type
        if filename.lower().endswith('.pdf'):
            content = await self._extract_pdf_content(file)
        elif filename.lower().endswith('.docx'):
            content = await self._extract_docx_content(file)
        elif filename.lower().endswith(('.txt', '.md')):
            content = file.read().decode('utf-8')
        else:
            raise ValueError(f"Unsupported file type: {filename}")
        
        # Extract title from filename
        title = filename.rsplit('.', 1)[0].replace('_', ' ').replace('-', ' ').title()
        
        return await self.ingest_document(
            title=title,
            content=content,
            document_type=document_type,
            original_filename=filename,
            **kwargs,
        )
    
    async def get_document(self, document_id: str) -> ERPDocument | None:
        """Get a document by ID."""
        result = await self.session.execute(
            select(ERPDocument).where(
                and_(
                    ERPDocument.id == document_id,
                    ERPDocument.tenant_id == self.tenant_id,
                )
            )
        )
        return result.scalar_one_or_none()
    
    async def search_documents(
        self,
        query: str,
        module: str | None = None,
        document_type: DocumentType | None = None,
        limit: int = 10,
    ) -> list[tuple[ERPDocumentChunk, float]]:
        """
        Search documents using semantic similarity.
        
        Args:
            query: Search query
            module: Filter by module
            document_type: Filter by document type
            limit: Maximum results
            
        Returns:
            List of (chunk, relevance_score) tuples
        """
        self.logger.log_operation_start(
            "search_documents",
            tenant_id=self.tenant_id,
            query=query[:50],
        )
        
        # Generate query embedding
        query_embedding = await ai_client.generate_single_embedding(query)
        
        # Search vector store
        vector_store = tenant_vector_store.get_store(self.tenant_id, "erp_docs")
        results = vector_store.search(
            query_embedding,
            top_k=limit * 2,  # Get more for filtering
            threshold=0.5,
        )
        
        # Fetch chunks and filter
        chunks_with_scores = []
        for result in results:
            chunk_result = await self.session.execute(
                select(ERPDocumentChunk, ERPDocument)
                .join(ERPDocument, ERPDocument.id == ERPDocumentChunk.document_id)
                .where(
                    and_(
                        ERPDocumentChunk.id == result.id,
                        ERPDocumentChunk.tenant_id == self.tenant_id,
                        ERPDocument.is_active == True,
                    )
                )
            )
            row = chunk_result.first()
            
            if row:
                chunk, doc = row
                
                # Apply filters
                if module and doc.module != module:
                    continue
                if document_type and doc.document_type != document_type:
                    continue
                
                chunks_with_scores.append((chunk, result.score))
                
                if len(chunks_with_scores) >= limit:
                    break
        
        self.logger.log_operation_complete(
            "search_documents",
            tenant_id=self.tenant_id,
            results_count=len(chunks_with_scores),
        )
        
        return chunks_with_scores
    
    async def list_documents(
        self,
        module: str | None = None,
        document_type: DocumentType | None = None,
        limit: int = 50,
        offset: int = 0,
    ) -> tuple[list[ERPDocument], int]:
        """List documents with filters."""
        config = await self._get_erp_config()
        if not config:
            return [], 0
        
        conditions = [
            ERPDocument.tenant_id == self.tenant_id,
            ERPDocument.erp_config_id == config.id,
            ERPDocument.is_active == True,
        ]
        
        if module:
            conditions.append(ERPDocument.module == module)
        
        if document_type:
            conditions.append(ERPDocument.document_type == document_type)
        
        # Count
        count_result = await self.session.execute(
            select(func.count(ERPDocument.id)).where(and_(*conditions))
        )
        total = count_result.scalar_one()
        
        # Get documents
        result = await self.session.execute(
            select(ERPDocument)
            .where(and_(*conditions))
            .order_by(ERPDocument.title)
            .limit(limit)
            .offset(offset)
        )
        documents = result.scalars().all()
        
        return list(documents), total
    
    async def update_document(
        self,
        document_id: str,
        **updates: Any,
    ) -> ERPDocument | None:
        """Update a document."""
        doc = await self.get_document(document_id)
        if not doc:
            return None
        
        content_changed = "content_raw" in updates
        
        for key, value in updates.items():
            if hasattr(doc, key) and key not in ["id", "tenant_id"]:
                setattr(doc, key, value)
        
        doc.updated_at = datetime.utcnow()
        
        # Re-process if content changed
        if content_changed:
            doc.content_cleaned = await self._clean_content(doc.content_raw)
            doc.content_summary = await self._generate_summary(doc.content_cleaned)
            doc.is_processed = False
            await self._process_document(doc)
        
        await self.session.flush()
        return doc
    
    async def delete_document(self, document_id: str) -> bool:
        """Delete a document (soft delete)."""
        doc = await self.get_document(document_id)
        if not doc:
            return False
        
        doc.is_active = False
        doc.updated_at = datetime.utcnow()
        
        # Remove from vector store
        vector_store = tenant_vector_store.get_store(self.tenant_id, "erp_docs")
        chunk_ids = await self._get_chunk_ids(document_id)
        vector_store.delete(chunk_ids)
        
        await self.session.flush()
        return True
    
    async def _process_document(self, document: ERPDocument) -> None:
        """Process document: chunk, embed, and index."""
        try:
            # Delete existing chunks
            await self.session.execute(
                ERPDocumentChunk.__table__.delete().where(
                    ERPDocumentChunk.document_id == document.id
                )
            )
            
            # Chunk content
            chunks_data = text_chunker.chunk_with_headers(document.content_cleaned)
            
            # Generate embeddings in batches
            chunk_texts = [c["content"] for c in chunks_data]
            embeddings = await ai_client.generate_embeddings(chunk_texts)
            
            # Create chunk records and index
            chunk_ids = []
            for i, (chunk_data, embedding) in enumerate(zip(chunks_data, embeddings)):
                chunk = ERPDocumentChunk(
                    id=str(uuid4()),
                    tenant_id=self.tenant_id,
                    document_id=document.id,
                    chunk_index=i,
                    content=chunk_data["content"],
                    token_count=len(chunk_data["content"].split()),
                    start_position=chunk_data.get("start_position"),
                    parent_headers=chunk_data.get("parent_headers", []),
                    embedding_model=settings.ai.embedding_model,
                    metadata={
                        "document_title": document.title,
                        "module": document.module,
                        "menu_path": document.erp_menu_path,
                    },
                )
                
                self.session.add(chunk)
                chunk_ids.append(chunk.id)
                
                # Add to vector store
                vector_store = tenant_vector_store.get_store(self.tenant_id, "erp_docs")
                vector_store.add_documents(
                    ids=[chunk.id],
                    embeddings=[embedding],
                    documents=[{
                        "content": chunk_data["content"],
                        "document_id": document.id,
                        "document_title": document.title,
                        "module": document.module,
                        "section_title": chunk_data.get("parent_headers", [""])[0] if chunk_data.get("parent_headers") else None,
                    }],
                )
            
            document.is_processed = True
            document.last_processed_at = datetime.utcnow()
            document.processing_error = None
            
        except Exception as e:
            self.logger.log_operation_failed(
                "process_document",
                e,
                tenant_id=self.tenant_id,
                document_id=document.id,
            )
            document.is_processed = False
            document.processing_error = str(e)
        
        await self.session.flush()
    
    async def _get_erp_config(self) -> ERPConfiguration | None:
        """Get ERP configuration for tenant."""
        result = await self.session.execute(
            select(ERPConfiguration).where(
                and_(
                    ERPConfiguration.tenant_id == self.tenant_id,
                    ERPConfiguration.is_active == True,
                )
            )
        )
        return result.scalar_one_or_none()
    
    async def _get_chunk_ids(self, document_id: str) -> list[str]:
        """Get all chunk IDs for a document."""
        result = await self.session.execute(
            select(ERPDocumentChunk.id).where(
                ERPDocumentChunk.document_id == document_id
            )
        )
        return [row[0] for row in result.all()]
    
    async def _clean_content(self, content: str) -> str:
        """Clean document content."""
        # Basic cleaning - remove excess whitespace
        import re
        
        cleaned = content
        cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)  # Multiple newlines
        cleaned = re.sub(r'[ \t]+', ' ', cleaned)  # Multiple spaces
        cleaned = cleaned.strip()
        
        return cleaned
    
    async def _generate_summary(self, content: str) -> str:
        """Generate document summary."""
        prompt = f"""
Summarize this ERP documentation in 2-3 sentences. Focus on:
- What process or feature it covers
- Key steps or information
- Who would use this

Content:
{content[:4000]}
"""
        
        return await ai_client.generate_text(
            prompt,
            system_prompt="Summarize ERP documentation concisely.",
            max_tokens=200,
        )
    
    async def _extract_keywords(self, content: str) -> list[str]:
        """Extract keywords from content."""
        prompt = f"""
Extract 5-10 keywords or key phrases from this ERP documentation.
Return as a JSON array of strings.

Content:
{content[:3000]}
"""
        
        result = await ai_client.generate_structured(
            prompt,
            {"type": "array", "items": {"type": "string"}},
        )
        
        return result if isinstance(result, list) else []
    
    async def _extract_pdf_content(self, file: BinaryIO) -> str:
        """Extract text from PDF file."""
        from pypdf import PdfReader
        
        reader = PdfReader(file)
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return "\n\n".join(text_parts)
    
    async def _extract_docx_content(self, file: BinaryIO) -> str:
        """Extract text from DOCX file."""
        from docx import Document
        
        doc = Document(file)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        return "\n\n".join(text_parts)
